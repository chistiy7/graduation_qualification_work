#!/usr/bin/env python3
"""Convert diploma_v1.md to formatted .docx (НИРС template)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips


FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADING = Pt(14)
FIRST_LINE_INDENT = Cm(1.25)
HEADING_LEFT_INDENT = Cm(2.0)

MARGINS = {
    "left": Cm(3.0),
    "right": Cm(1.0),
    "top": Cm(2.0),
    "bottom": Cm(3.0),
}


def set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = MARGINS["left"]
    section.right_margin = MARGINS["right"]
    section.top_margin = MARGINS["top"]
    section.bottom_margin = MARGINS["bottom"]


def add_body_paragraph(doc: Document, text: str, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = p.add_run(text.strip())
    set_run_font(run)
    return p


def add_heading_paragraph(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = HEADING_LEFT_INDENT if level >= 1 else Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.strip())
    set_run_font(run, size=FONT_SIZE_HEADING, bold=True)
    return p


def add_center_paragraph(doc: Document, text: str, *, bold=False, size=FONT_SIZE_BODY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text.strip())
    set_run_font(run, size=size, bold=bold)
    return p


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        if re.match(r"^\|\s*[-:]+\s*\|", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = parse_table_rows(table_lines)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(cell_text)
            set_run_font(run, bold=(i == 0))


def add_numbered_item(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text.strip())
    set_run_font(run)


def add_bullet_item(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text.strip())
    set_run_font(run)


def add_rich_paragraph(doc: Document, text: str, *, indent=True):
    """Paragraph with **bold** segments."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)
    lines = raw.splitlines()

    doc = Document()
    configure_document(doc)

    i = 0
    title_mode = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_page_break(doc)
            i += 1
            title_mode = False
            continue

        if stripped.startswith("# "):
            add_heading_paragraph(doc, stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading_paragraph(doc, stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading_paragraph(doc, stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_numbered_item(doc, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            add_bullet_item(doc, stripped[2:])
            i += 1
            continue

        if title_mode and (
            stripped.startswith("Министерство")
            or stripped.startswith("Федеральное")
            or stripped.startswith("«")
            or stripped.startswith("(МГТУ")
            or stripped.startswith("ФАКУЛЬТЕТ")
            or stripped.startswith("КАФЕДРА")
            or stripped.startswith("РАСЧЁТНО")
            or stripped.startswith("к выпускной")
            or stripped.startswith("Студент")
            or stripped.startswith("Руководитель")
            or stripped.startswith("2026")
            or stripped.startswith("УТВЕРЖДАЮ")
            or stripped.startswith("Заведующий")
            or stripped.startswith("ЗАДАНИЕ")
            or stripped.startswith("на выполнение")
            or stripped.startswith("Тема работы")
            or stripped.startswith("Техническое задание")
            or stripped.startswith("Оформление")
            or stripped.startswith("Дата выдачи")
            or stripped.startswith("«___»")
            or stripped.startswith("__________________")
            or stripped.startswith("(подпись")
            or stripped.startswith("(фамилия")
            or stripped == "Содержание"
        ):
            if stripped == "Содержание":
                add_center_paragraph(doc, stripped, bold=True, size=FONT_SIZE_HEADING)
            elif stripped.startswith("РАСЧЁТНО") or stripped.startswith("ЗАДАНИЕ"):
                add_center_paragraph(doc, stripped, bold=True, size=FONT_SIZE_HEADING)
            elif stripped.startswith("«") and "Разработка" in stripped:
                add_center_paragraph(doc, stripped, bold=False)
            else:
                add_center_paragraph(doc, stripped)
            i += 1
            continue

        if stripped == "Содержание":
            add_center_paragraph(doc, stripped, bold=True, size=FONT_SIZE_HEADING)
            i += 1
            continue

        if re.match(r"^\d+\.\d+\.", stripped) or stripped.startswith("Глава "):
            add_body_paragraph(doc, stripped, indent=False)
            i += 1
            continue

        if "**" in stripped:
            add_rich_paragraph(doc, stripped)
        else:
            add_body_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


def main():
    base = Path(__file__).resolve().parent.parent
    md_path = base / "diploma_v1.md"
    docx_path = base / "diploma_v1.docx"
    if len(sys.argv) > 1:
        md_path = Path(sys.argv[1])
    if len(sys.argv) > 2:
        docx_path = Path(sys.argv[2])
    convert_md_to_docx(md_path, docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    main()
